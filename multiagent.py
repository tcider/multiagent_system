# Модуль для работы с pdf
# pip install pymupdf
import fitz

# Модуль для работы с Word
# pip install python-docx
from docx import Document
from docx.shared import Inches

# Конверте doc в pdf
# pip install docx2pdf
from docx2pdf import convert

# Для копирования файлов
import shutil

# Для работы с фалами в ОС
import os

# Задаем пути к папке источнику и папке с результатом
# Путь к папке с входящими файлами
SOURCE_PATH = "source"
# Путь к папке с результатами
RESULT_PATH = "result"
# Путь к папке с найденными картинками
IMAGE_PATH  = "images"
# Путь к папке с отдельными текстами
DOCS_PATH  = "docs"


# Вспомогательная функция очистики имени файла от расширения
def clear(file_name):
    file_name_list = file_name.split('.')
    file_name_list.pop(len(file_name_list) - 1)
    return ".".join(file_name_list)


############### 1 тип Агентов исполнителей - Агент заготовитель ####################
# Класс агентов возвращающих список заданного типа файлов
class   AgentFileSearch:
    # Списки найденных файлов
    file_all_list = []
    # МНожестов для хранения doc файлов
    doc_files = set()

    # При инициализации класса агент исследует входящий катлог на типы айлов и составлет их списки
    def __init__(self):
        print("Агент поиска файлов: Инициализирован...")
        # Тк word сохраняет каритинки внтури в своем формате для их извлечения конвертируем в pdf
        for root, dirs, files in os.walk(SOURCE_PATH):
            for file in files:
                if file.endswith("docx") or file.endswith("DOCX"):
                    file_path = os.path.join(root, file)
                    convert(file_path)
                    self.doc_files.add(clear(file))
                    print()
        for root, dirs, files in os.walk(SOURCE_PATH):
            for file in files:
                print("Агент поиска файлов: Файл", file, "найден...")
                file_path = os.path.join(root, file)
                self.file_all_list.append((file_path, file))


############### 2 тип Агентов исполнителей - Агенты обработчики файлов ####################
# Класс агента обрабатывающего Pdf файл (и Word так же тк они конвертируются через pdf)
class   AgentPdf:
    # Имя афйла Pdf
    file_name = ""
    # Путь к файлу
    file = ""
    # Для хранения обьекта документа Pdf
    doc = None
    # id агента
    id = -1
    # Словарь для хранения метаданных файла
    meta = ""
    # Кол-во страниц в документе
    page_num = 0
    # Кол-во картинок в документе
    image_num = 0
    # doc файлы
    doc_files = set()

    def __init__(self, file_path, file, agent_id, doc_files):
        self.file = file_path
        self.file_name = file
        self.id = agent_id
        self.doc_files = doc_files
        self.doc = fitz.open(file_path)
        self.meta = self.get_header(self.doc.metadata)
        self.page_num = self.doc.pageCount

    # Метод соаздющий текстовое описание матаданных документа
    def get_header(self, voc):
        res = ""
        for key, value in voc.items():
            if key == "format" and clear(self.file_name) in self.doc_files:
                value = "DOCX"
            res += str(key) + ": " + str(value) + ", "
        return res

    # Метод возвращает текст из pdf Документа
    def get_text(self):
        res = ""
        for current_page in range(len(self.doc)):
            page = self.doc.loadPage(current_page)
            page_text = page.getText("text")
            res += page_text
        text_agent = AgentOneTextResult(self.file_name)
        text_agent.write(self.file_name, self.meta, self.page_num, res)
        text_agent.save()
        return res

    # Метод получения картинок из pdf
    def get_images(self):
        page_count = 0
        res = []
        for i in range(len(self.doc)):
            for img in self.doc.getPageImageList(i):
                xref = img[0]
                pix = fitz.Pixmap(self.doc, xref)
                pix1 = fitz.Pixmap(fitz.csRGB, pix)
                page_count += 1
                file_name = "/%s_img_%s(page_%s).png" % (self.file_name, page_count, i+1)
                file_type = "png"
                file_path = RESULT_PATH + "/" + IMAGE_PATH + file_name
                pix1.writePNG(file_path)
                pix1 = None
                res.append((file_path, file_type))
        self.image_num = page_count
        return res


# Агент обработки обычных текстов
class   AgentTxt:
    # Имя файла
    file_name = ""
    # Путь к файлу
    file = ""
    # id агента
    id = -1
    # Словарь для хранения метаданных файла
    meta = ""
    # Кол-во страниц в документе
    page_num = 0
    # Кол-во картинок в документе
    image_num = 0

    def __init__(self, file_path, file, agent_id):
        self.file = file_path
        self.file_name = file
        self.id = agent_id
        self.meta = "Plain text file."

    # Метод возвращает текст
    def get_text(self):
        f = open(self.file, encoding="cp1251")
        res = f.read()
        f.close()
        text_agent = AgentOneTextResult(self.file_name)
        text_agent.write(self.file_name, self.meta, self.page_num, res)
        text_agent.save()
        return res

    # Метод получения картинок
    def get_images(self):
        return []


# Агент обработки картинок
class   AgentImg:
    # Имя файла
    file_name = ""
    # Путь к файлу
    file = ""
    # id агента
    id = -1
    # Словарь для хранения метаданных файла
    meta = ""
    # Кол-во страниц в документе
    page_num = 0
    # Кол-во картинок в документе
    image_num = 0

    def __init__(self, file_path, file, agent_id):
        self.file = file_path
        self.file_name = file
        self.id = agent_id
        self.meta = "Image file."

    # Метод возвращает текст
    def get_text(self):
        return []

    # Метод получения картинок
    def get_images(self):
        res = []
        self.image_num = 1
        ext_list = self.file_name.split('.')
        file_type = ext_list[len(ext_list) - 1]
        file_name = clear(self.file_name) + "_img" + "." + file_type
        file_path = RESULT_PATH + "/" + IMAGE_PATH + "/" + file_name
        shutil.copyfile(self.file, file_path)
        res.append((file_path, file_type))
        return res


############### 3 тип Агентов исполнителей - Агенты записывающие результат ####################
# Агент записи текста в один общий файл
class   AgentTextResult:
    # Имя файла с результатами
    text_file_name = "all_text.doc"
    # Обьект Word файла
    text_doc = None
    # Сообщение вывода сохраняю в перменную тк класс будет родителем др класса где будет др сообщение
    msg1 = "Агент записи текстов общий: Запись текста в общий файл из файла %s..."
    msg2 = "Агент записи текстов общий: Сохранение общего файла..."
    sub_path = "/"


    def __init__(self):
        print("Агент записи текстов общий: Иницализирован...")
        self.text_doc = Document()

    # Запись текста в буфер
    def write(self, file_name, meta, page_num, text):
        print(self.msg1 % clear(file_name))
        self.text_doc.add_heading(clear(file_name), level=1)
        self.text_doc.add_heading("Мета теги", level=2)
        self.text_doc.add_paragraph(meta)
        self.text_doc.add_heading("Страниц в тексте - %s" % page_num, level=2)
        self.text_doc.add_heading("Текстовое содержимое", level=2)
        self.text_doc.add_paragraph(text)
        self.text_doc.add_page_break()

    # Сохранение буфера в файл
    def save(self):
        print(self.msg2)
        self.text_doc.save(RESULT_PATH + self.sub_path + self.text_file_name)


# Линейный агент записи текстов, унаследованный от старшего агента
class   AgentOneTextResult(AgentTextResult):
    msg1 = "Агент записи текстов: Запись текста в индивидуальный файл из файла %s..."
    msg2 = "Агент записи текстов: Сохранение индивидуального файла..."
    sub_path = "/"+ DOCS_PATH +"/"

    def __init__(self, file_name):
        print("Агент записи текстов: Иницализирован...")
        self.text_doc = Document()
        self.text_file_name = clear(file_name) + "_text.doc"


# Запись картинок
class   AgentImageResult:
    # Имя файла с результатами
    img_file_name = "all_images.doc"
    # Обьект Word файла
    img_doc = None

    def __init__(self):
        print("Агент записи изображений: Иницализирован...")
        self.img_doc = Document()

    # Запись картинок в буфер
    def write(self, file_name, img_num, img_list):
        print("Агент записи изображений: Запись изображений из файла %s..." % clear(file_name))
        self.img_doc.add_heading(clear(file_name), level=1)
        self.img_doc.add_heading("Изображений в файле - %s" % img_num, level=2)
        for file, type in img_list:
            self.img_doc.add_heading("Формат - %s" % type, level=2)
            self.img_doc.add_picture(file, width=Inches(5))
        self.img_doc.add_page_break()

    # Сохранение буфера в файл
    def save(self):
        print("Агент записи изображений: Сохранение итогов в файл...")
        self.img_doc.save(RESULT_PATH + "/" + self.img_file_name)


############### 4 тип самый главный Агент - Лидер ####################
# Класс агента лидера (управляющего другими агентами)
class   AgentLeader:
    stat = []
    # Обьекта агента поиска файлов
    agent_file_search = None
    # Спиок для хранения задач(файлов для обработки) и id агентов кому задача поручена
    file_job_list = []
    # Список обьектов - агентов исполнителей, обработчиков файлов
    agent_list = []
    # Обьект Аегнта записи текстовго результата
    agent_text = None
    # Обьект Аегнта записи картинок
    agent_img = None
    # word файлы
    doc_files = set()

    # При инициализации сразу создается агент поисковик файлов
    def __init__(self):
        print("Агент лидер: Инициализирован...")
        agent_file_search = AgentFileSearch()
        self.doc_files = agent_file_search.doc_files
        # Заполняем список задач пока -1 id агентов кому задача будет дана
        for file_path, file in agent_file_search.file_all_list:
            self.file_job_list.append([file_path, file, -1])

    # Метод раздачи работы агентам исполнителям по списку данному агентом поисковиком
    def job_distribute(self):
        print("Агент лидер: Запущено распределение задач среди Агентов исполнителей...")
        agent_id = 0
        for i in range(len(self.file_job_list)):
            file_path = self.file_job_list[i][0]
            file = self.file_job_list[i][1]
            if file_path.endswith(".pdf") or file_path.endswith(".PDF"):
                agent = AgentPdf(file_path, file, agent_id, self.doc_files)
                self.agent_list.append(agent)
                self.file_job_list[i][2] = agent_id
                agent_id += 1
            elif file_path.endswith(".txt") or file_path.endswith(".TXT"):
                agent = AgentTxt(file_path, file, agent_id)
                self.agent_list.append(agent)
                self.file_job_list[i][2] = agent_id
                agent_id += 1
            elif file_path.endswith(".jpg") or file_path.endswith(".JPG") or \
                    file_path.endswith(".png") or file_path.endswith(".PNG"):
                agent = AgentImg(file_path, file, agent_id)
                self.agent_list.append(agent)
                self.file_job_list[i][2] = agent_id
                agent_id += 1

    # Инициализация Агентов записи реузльтатов в файлы
    def create_res_agents(self):
        self.agent_text = AgentTextResult()
        self.agent_img = AgentImageResult()

    # Метод получния результатов работы от агентов исполнитлей и передачи ее агентам записи результатов
    def job_done(self):
        for i in range(len(self.agent_list)):
            text = self.agent_list[i].get_text()
            meta = self.agent_list[i].meta
            page_num = self.agent_list[i].page_num
            file_name = self.agent_list[i].file_name
            agent_id = self.agent_list[i].id
            print("Агент исполнитель №%s: Обработка файла %s..." % (agent_id, file_name))
            if (len(text)):
                self.agent_text.write(file_name, meta, page_num, text)
            # Обработка изображений
            image_list = self.agent_list[i].get_images()
            image_num = self.agent_list[i].image_num
            if image_num:
                self.agent_img.write(file_name, image_num, image_list)

    def close_res_files(self):
        self.agent_text.save()
        self.agent_img.save()

    def free_agents(self):
        print("Агент лидер: Очистка памяти от Агентов исполнителей...")
        for elem in self.agent_list:
            del elem


# Функция main выполняет роль командного центра (системы упраления агентами)
def main():
    # Создаем самого главного агента
    leader = AgentLeader()
    # Агент лидер раздает задачи
    leader.job_distribute()
    # Агент лидер инициализиурет агентов записи результатов
    leader.create_res_agents()
    # Агент лидер собирает результаты
    leader.job_done()
    # Агент лижер закрывает Агентов записи результата
    leader.close_res_files()
    # Агент лидер освобождает память от исполнителей)
    leader.free_agents()


if __name__ == "__main__":
    main()